from datetime import datetime
from docx import Document
from docx.shared import Inches
import pymysql


class DbInfo(object):
    host = ''
    user = ''
    password = ''
    database = ''
    charset = ''
    conn = None

    def __init__(self, host=None, user=None, password=None, database=None, charset=None):
        self.host = host
        self.user = user
        self.password = password
        self.database = database
        self.charset = charset

    def connMysql(self):
        self.conn = pymysql.connect(
            host=self.host, 
            user=self.user, 
            password=self.password, 
            database=self.database,
            charset=self.charset
            )
        return self.conn

    def closesql(self):
        self.conn.close()

    def execute(self, sql=None):
        cur = self.connMysql()
        with cur.cursor() as cursor:
            sqllist = sql
            cursor.execute(sqllist)
            result = cursor.fetchall()
            self.conn.commit()
        return result

    def getSchemalName(self, ConnName):
        '''根据数据库链接名获取此连接下的数据库名称信息'''
        sql = f'''
        select  BaseName  from core_sqldictionary where Name='{ConnName}'  group by BaseName
        '''
        return self.execute(sql)

    def getTableName(self, ConnName, SchemalName, TableName=None):
        '''根据数据库链接名 和数据库名 获取此连接下的schemal的表名称信息'''
        if TableName:
            sql = f'''
            select  `TableName`  \
            from core_sqldictionary where Name='{ConnName}' and \
            BaseName ='{SchemalName}' and TableName = '{TableName}' group by `TableName`
            '''
            return self.execute(sql)
        else:
            sql = f"select  `TableName`  from core_sqldictionary \
            where Name='{ConnName}' and BaseName ='{SchemalName}' and TableName = '{TableName}' \
            group by TableName limit 1"
            return self.execute(sql)

    def getTableInfo(self, ConnName, SchemalName=None, TableName=None):
        sql = f"""select `Name`, `BaseName`, `TableName`, `TableComment`, \
        `Field`, `Type`, `Extra` from core_sqldictionary \
        where Name = '{ConnName}' and BaseName='{SchemalName}' and TableName='{TableName}' """
        return self.execute(sql)


class ToWord:
    document = None

    def __init__(self, Host=None, User=None, Password=None, Database=None, Charset=None):
        self.turnOjb = DbInfo(
            host=Host, 
            user=User, 
            password=Password, 
            database=Database, 
            charset=Charset
            )
        self.createDoc()

    def createDoc(self):
        self.document = Document()
        self.document.add_heading('数据库文档', 0)
        self.p = self.document.add_paragraph('Yearning')
        self.p.add_run('Yearning').bold = True
        self.p = self.document.add_paragraph('导出日期: %s' % datetime.now())
        self.document.add_picture('libs/logo.png', width=Inches(8))

    def exportTables(self, Conn=None, Schemal=None, TableList=None):
        '''导出指定的一些表，TableList 为表名称列表 '''

        for tableName in TableList:
            tabSet = self.turnOjb.getTableName(
                ConnName=Conn,
                SchemalName=Schemal,
                TableName=tableName)
            self.document.add_page_break()
            self.document.add_heading(
                '%s' %[TB[0] for TB in tabSet][0], level=2
                )
            table = self.document.add_table(rows=1, cols=5)
            table.style = 'LightShading-Accent1'
            table.rows[0].cells[0].text = '字段名'
            table.rows[0].cells[1].text = '类型'
            table.rows[0].cells[2].text = '备注'
            columnSet = self.turnOjb.getTableInfo(ConnName=Conn, SchemalName=Schemal,
                                                  TableName='%s' % [TB[0] for TB in tabSet][0])
            for index, column in enumerate(columnSet):
                cells = table.add_row().cells
                cells[0].text = '%s' % column[4]
                cells[1].text = '%s' % column[5]
                cells[2].text = '%s' % column[6]
        time = datetime.now()
        self.document.save('./exportData/%s_%s_Dictionary_%s.docx' % (Conn, Schemal, time))
        return time

    def exportSchemal(self, Conn=None, Schemal=None):
        tabSet = self.turnOjb.getTableName(ConnName=Conn, SchemalName=Schemal)
        for tableName in tabSet:
            self.document.add_page_break()
            self.document.add_heading('%s : %s' % (tableName[0], tableName[1]), level=2)
            table = self.document.add_table(rows=1, cols=5)
            table.rows[0].cells[0].text = '字段名'
            table.rows[0].cells[1].text = '类型'
            table.rows[0].cells[2].text = '是否可以为空'
            table.rows[0].cells[3].text = '默认值'
            table.rows[0].cells[4].text = '备注'
            columnSet = self.turnOjb.getTableInfo(
                ConnName=Conn, 
                SchemalName=Schemal, 
                TableName='%s' % tableName[0])
            for index, column in enumerate(columnSet):
                cells = table.add_row().cells
                cells[0].text = '%s' % column[4]
                cells[1].text = '%s' % column[5]
                cells[2].text = '%s' % column[6]
                cells[3].text = '%s' % column[7]
                cells[4].text = '%s' % column[8]
        self.document.save('./exportData/%s_%s_数据字典.docx' % (Conn, Schemal))